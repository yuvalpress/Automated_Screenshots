import pyautogui #Screenshot utility

#Wev browser driver utilities
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from time import sleep #Sleep functionality

#Choose files and folders interface
from tkinter import filedialog
from tkinter import Tk

import os #Change % create dirs

import xlrd #Read excel file utility

import shutil #For folder exsitense check & for file compressing

import sys #For sau servers validation

import subprocess #For Sau servers validation

from docx import Document #For auto written docx files
from docx.shared import Inches

import logging #For logging

import requests


def shot(folder, name): #Screenshot functionality
    try:
        if os.path.isdir(folder):
            myScreenshot = pyautogui.screenshot()
            myScreenshot.save(r"{}\\{}.png".format(folder, name))
            logging.info("Saved {}.png to {}".format(name, folder))

        else:
            os.mkdir(folder)
            myScreenshot = pyautogui.screenshot()
            myScreenshot.save(r"{}\\{}.png".format(folder, name))
            logging.info("Saved {}.png to {}".format(name, folder))

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

def ping(ip, name):

    #ping requested server
    try:
        ping = subprocess.Popen(
            ["ping", "-n", "1", '{}'.format(ip)],
            stdout = subprocess.PIPE,
            stderr = subprocess.PIPE)

        out, error = ping.communicate()
    except:
        logging.error(error)
    
    

    print(str(out))

    if "TTL" in str(out):
        return True
    else: return False

def docxReady(ex, folder, tf):

    #Search excel file for Engineer data
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Engineer_Info")
        eng = {}
        for i in range(0, sheet.ncols, +1):
            eng[sheet.cell_value(0,i)] = sheet.cell_value(1,i)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from Engineer_Info tab has been retrieved Successfully")

    for server, answer in tf.items():
        if answer == True:

            #Read from excel file
            try:
                wb = xlrd.open_workbook(ex)
                sheet = wb.sheet_by_name(server)
                dic = {}
                idr = {}
                ops = {}
                for i in range(1, sheet.nrows, +1):
                    dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
                
                for i in range(1, sheet.nrows, +1):
                    idr[sheet.cell_value(i,1)] = sheet.cell_value(i,2)

                for i in range(1, sheet.nrows, +1):
                    ops[sheet.cell_value(i,1)] = sheet.cell_value(i,3)
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)

            logging.info("Creation of 3 data Dictionaries has been completed")

            #Create directory var
            serversMainFolder = '{}\\{} Validations'.format(folder, server)
            logging.info("Variable for {} server main validation directory has been created".format(server))

            if server != "3GI" and server != "Stargate":
                for key, value in dic.items():

                    #Checks if server is pingable
                    if ping(value, key):

                        logging.info("{} is pingable".format(key))

                        #Set current server folder
                        server_folder = "{}\\{}".format(serversMainFolder, key)
                        logging.info("Variable for server {} directory has been modified".format(key))
                        
                        #create docx file using template and "save" function
                        doc = Document(r"Docx Templates\\ESXi.docx")
                        logging.info("Docx file template for {} validation has been copied to a new file".format(server))

                        #Start modifieng the main table

                        try:
                            for row in doc.tables[0].rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        inline = p.runs
                                        for i in range(len(inline)):

                                            #Edit lines inside table
                                            if inline[i].text == "EngineerName": inline[i].text = eng.get('Engineer_Name')

                                            if inline[i].text == "CIRName": inline[i].text = eng.get('CIR_Name')
                                                
                                            if inline[i].text == "ProjectName": inline[i].text = eng.get('Project_Name')

                                            if inline[i].text == "RevisionNumber": inline[i].text = str(eng.get('Revision_Number'))

                                            if inline[i].text == "dateSequence": inline[i].text = str(eng.get('Date'))

                                            if inline[i].text == "ServerName": inline[i].text = key

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)
                        
                        logging.info("Main table has been modified according to {} server data".format(key))

                        try: 
                            for p in doc.paragraphs:
                                inline = p.runs
                                for i in range(len(inline)):

                                    #Start edit lines outside table
                                    if inline[i].text == " IPv4": inline[i].text = " {}".format(idr.get(key))
                                    if inline[i].text == "SystemVar": inline[i].text = " {}".format(ops.get(key))
                            logging.info("Idrac & OS data inside Docx document has been modified accoeding to {} server data".format(key))

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)

                        #Add screenshots here
                        try:
                            for screenshot in os.listdir(server_folder):
                                if "png" in screenshot or "jpg" in screenshot:
                                    doc.add_picture("{}\\{}".format(server_folder, screenshot), width=Inches(7.0))
                            logging.info("Screenshots from {} directory has been added to the docx document successfully".format(key))

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)

                        
                        #Save docx file modifications
                        doc.save("{}\\{}.docx".format(server_folder, key))
                        logging.info("Docx document for {} server was saved successfully".format(key))
                    
                    elif ping(idr.get(key), key):

                        logging.info("{} iDrac is pingable".format(key))

                        #Set current server folder
                        server_folder = "{}\\{}".format(serversMainFolder, key)
                        logging.info("Variable for {}'s iDrac directory has been modified".format(key))
                        
                        #create docx file using template and "save" function
                        doc = Document(r"Docx Templates\\{}.docx".format(server))
                        logging.info("Docx file template for {} validation has been copied to a new file".format(server))

                        #Start modifieng the main table

                        try:
                            for row in doc.tables[0].rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        inline = p.runs
                                        for i in range(len(inline)):

                                            #Edit lines inside table
                                            if inline[i].text == "EngineerName": inline[i].text = eng.get('Engineer_Name')

                                            if inline[i].text == "CIRName": inline[i].text = eng.get('CIR_Name')
                                                
                                            if inline[i].text == "ProjectName": inline[i].text = eng.get('Project_Name')

                                            if inline[i].text == "RevisionNumber": inline[i].text = str(eng.get('Revision_Number'))

                                            if inline[i].text == "dateSequence": inline[i].text = str(eng.get('Date'))

                                            if inline[i].text == "ServerName": inline[i].text = key

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)
                        
                        logging.info("Main table has been modified according to {} server data".format(key))

                        try: 
                            for p in doc.paragraphs:
                                inline = p.runs
                                for i in range(len(inline)):

                                    #Start edit lines outside table
                                    if inline[i].text == " IPv4": inline[i].text = " {}".format(idr.get(key))
                                    if inline[i].text == "SystemVar": inline[i].text = " {}".format(ops.get(key))
                            logging.info("Idrac & OS data inside Docx document has been modified accoeding to {} server data".format(key))

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)

                        #Add screenshots here
                        try:
                            for screenshot in os.listdir(server_folder):
                                if "png" in screenshot or "jpg" in screenshot:
                                    doc.add_picture("{}\\{}".format(server_folder, screenshot), width=Inches(7.0))
                            logging.info("Screenshots from {} directory has been added to the docx document successfully".format(key))

                        except (RuntimeError, TypeError, NameError) as e:
                            logging.error(e)

                        
                        #Save docx file modifications
                        doc.save("{}\\{}.docx".format(server_folder, key))
                        logging.info("Docx document for {} server was saved successfully".format(key))

                    
                    else: logging.warning("Docx will no be created for server {} for it is not pingable".format(key))
            
            else:

                logging.info("Starting {} docx file creation".format(server))
                
                #create docx file using template and "save" function
                doc = Document(r"Docx Templates\\{}.docx".format(server))
                logging.info("Docx file template for {} validation has been copied to a new file".format(server))

                #Start modifieng the main table

                try:
                    for row in doc.tables[0].rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                inline = p.runs
                                for i in range(len(inline)):

                                    #Edit lines inside table
                                    if inline[i].text == "EngineerName": inline[i].text = eng.get('Engineer_Name')

                                    if inline[i].text == "CIRName": inline[i].text = eng.get('CIR_Name')
                                        
                                    if inline[i].text == "ProjectName": inline[i].text = eng.get('Project_Name')

                                    if inline[i].text == "RevisionNumber": inline[i].text = str(eng.get('Revision_Number'))

                                    if inline[i].text == "dateSequence": inline[i].text = str(eng.get('Date'))

                except (RuntimeError, TypeError, NameError) as e:
                    logging.error(e)
                
                logging.info("Main table has been modified according to {} server data".format(server))

                #Add screenshots here
                try:
                    for screenshot in os.listdir(serversMainFolder):
                        print(screenshot)
                        if "png" in screenshot or "jpg" in screenshot:
                            doc.add_picture("{}\\{}".format(serversMainFolder, screenshot), width=Inches(7.0))
                    logging.info("Screenshots from {} directory has been added to the docx document successfully".format(server))
                except (RuntimeError, TypeError, NameError) as e:
                    logging.error(e)
  
                #Save docx file modifications
                doc.save("{}\\{}-Checklist.docx".format(serversMainFolder, server))
                logging.info("Docx document for {} server was saved successfully".format(server))    
                    
def tbos(ex, folder): #Tbos servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Tbos")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from Tbos tab has been retrieved Successfully")

    #Create Tbos directory or not
    try:
        if os.path.isdir('{}\\Tbos Validations'.format(folder)):
            shutil.rmtree('{}\\Tbos Validations'.format(folder))

        folder = "{}\\Tbos Validations".format(folder)
        os.mkdir(folder)
        logging.info("Tbos Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Tbos validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():
        
        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create a server directory
            if os.path.isdir('{}\\{}'.format(folder, key)):
                shutil.rmtree('{}\\{}'.format(folder, key))

            folder1 = "{}\\{}".format(folder, key)
            os.mkdir(folder1)
            logging.info("A folder for server {} has been created at {}".format(key, folder1))
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Tbos {} '{}' {}".format(userPath, value, folder1, key)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
    logging.info("Executed iDrac validation for Tbos servers")
    iDRAC(ex, folder, "Tbos")

def kafka(ex, folder): #Kafka servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Kafka")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    logging.info("Data from Kafka tab has been retrieved Successfully")

    #Create Kafka directory or not
    try:
        if os.path.isdir('{}\\Kafka Validations'.format(folder)):
            shutil.rmtree('{}\\Kafka Validations'.format(folder))

        folder = "{}\\Kafka Validations".format(folder)
        os.mkdir(folder)
        logging.info("Kafka Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Kafka validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():

        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create a server directory
            if os.path.isdir('{}\\{}'.format(folder, key)):
                shutil.rmtree('{}\\{}'.format(folder, key))

            folder1 = "{}\\{}".format(folder, key)
            os.mkdir(folder1)
            logging.info("A folder for server {} has been created at {}".format(key, folder1))
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Kafka {} '{}' {}".format(userPath, value, folder1, key)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
    logging.info("Executed iDrac validation for Kafka servers")
    iDRAC(ex, folder, "Kafka")

def mde(ex, folder): #MDE servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("MDE")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    logging.info("Data from MDE tab has been retrieved Successfully")

    #Create MDE directory or not
    try:
        if os.path.isdir('{}\\MDE Validations'.format(folder)):
            shutil.rmtree('{}\\MDE Validations'.format(folder))

        folder = "{}\\MDE Validations".format(folder)
        os.mkdir(folder)
        logging.info("MDE Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #MDE validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():
        #Checks if server is pingable
        if ping(value, key):
            
            logging.info("{} is pingable".format(key))

            #Create a server directory
            if os.path.isdir('{}\\{}'.format(folder, key)):
                shutil.rmtree('{}\\{}'.format(folder, key))

            folder1 = "{}\\{}".format(folder, key)
            os.mkdir(folder1)
            logging.info("A folder for server {} has been created at {}".format(key, folder1))
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' MDE {} '{}' {}".format(userPath, value, folder1, key)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
        
    logging.info("Executed iDrac validation for MDE servers")
    #iDRAC(ex, folder, "MDE")

def presto(ex, folder): #Presto servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Presto")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from Presto tab has been retrieved Successfully")

    #Create Presto directory or not
    try:
        if os.path.isdir('{}\\Presto Validations'.format(folder)):
            shutil.rmtree('{}\\Presto Validations'.format(folder))

        folder = "{}\\Presto Validations".format(folder)
        os.mkdir(folder)
        logging.info("Presto Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Presto validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():

        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create a server directory
            if os.path.isdir('{}\\{}'.format(folder, key)):
                shutil.rmtree('{}\\{}'.format(folder, key))

            folder1 = "{}\\{}".format(folder, key)
            os.mkdir(folder1)
            logging.info("A folder for server {} has been created at {}".format(key, folder1))
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Presto {} '{}' {}".format(userPath, value, folder1, key)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
    logging.info("Executed iDrac validation for Presto servers")
    iDRAC(ex, folder, "Presto")

def workers(ex, folder): #Workers servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Workers")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from Workers tab has been retrieved Successfully")

    #Create Workers directory or not
    try:
        if os.path.isdir('{}\\Workers Validations'.format(folder)):
            shutil.rmtree('{}\\Workers Validations'.format(folder))

        folder = "{}\\Workers Validations".format(folder)
        os.mkdir(folder)
        logging.info("Workers Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Workers validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():

        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create a server directory
            if os.path.isdir('{}\\{}'.format(folder, key)):
                shutil.rmtree('{}\\{}'.format(folder, key))

            folder1 = "{}\\{}".format(folder, key)
            os.mkdir(folder1)
            logging.info("A folder for server {} has been created at {}".format(key, folder1))
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Workers {} '{}' {}".format(userPath, value, folder1, key)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
    #iDRAC(ex, folder, "Workers")
    logging.info("Executed iDrac validation for worker servers")

def threegi(ex, folder): #3GI servers screenshot creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("3GI")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    # Get mps server value
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("3GI")
        mps = {}
        for i in range(1, sheet.nrows, +1):
            mps[sheet.cell_value(i,1)] = sheet.cell_value(i,4)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from 3GI tab has been retrieved Successfully")

    #Create 3GI directory or not
    try:
        if os.path.isdir('{}\\3GI Validations'.format(folder)):
            shutil.rmtree('{}\\3GI Validations'.format(folder))

        folder = "{}\\3GI Validations".format(folder)
        os.mkdir(folder)
        logging.info("3GI Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #3GI validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    #Set security certificate as trusted
    options = webdriver.ChromeOptions()
    options.add_argument('ignore-certificate-errors')
    options.add_experimental_option("excludeSwitches", ['enable-automation']) #set banner of controlled chrome off
    logging.info("Security certificate has been set Successfuly")

    #create browser element + assign chrome driver
    browser = webdriver.Chrome('Files\\chromedriver.exe', chrome_options=options)
    browser.implicitly_wait(90)

    #Set wait element time
    wait = WebDriverWait(browser, 90)

    for key, value in dic.items():
        #Checks if server is pingable
        print(value)
        if ping(value, key):
            
            logging.info("{} is pingable".format(key))
            
            #Execute remote powershell script
            if mps.get(key):
                try:
                    userPath = os.environ['USERPROFILE']
                    p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' 3GI {} '{}' {}".format(userPath, value, folder, key)])
                    p.communicate()
                    logging.info("run_on_remote.ps1 script has been executed successfully")

                except (RuntimeError, TypeError, NameError) as e:
                    logging.error(e)
            
            if mps.get(key) == 0:
                try:

                    browser.get("https://{}/ui/#/host/vms".format(value)) #open browser window

                    #login into ESXi web page
                    ESXi_username = browser.find_element_by_id('username')
                    ESXi_password = browser.find_element_by_id('password')
                    ESXi_username.send_keys("root")
                    ESXi_password.send_keys("password")
                    browser.find_element_by_id("submit").click()

                    logging.info("Logged into ESXi web interface Successfully")

                    browser.maximize_window() #maximize browser window
                    wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='createVMButton']")))

                    #Take a screenshot of all the vms
                    shot(folder, "vms")
                    logging.info("Captured a screenshot of Virtual Machines page")

                    #Take a screenshot of the network span
                    browser.find_element_by_xpath("/html/body/div/div/div[1]/div/div[2]/div/div[2]/div[1]/div/div[2]/div/div/ul/li[4]/div/a/span[2]").click()
                    wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='portgroupGrid']/div/div[2]/div/table/thead/tr/th[1]")))
                    sleep(2)
                    shot(folder,"network")
                    logging.info("Captured a screenshot of Network page")

                    #Take a screenshot of the Licensing span
                    browser.find_element_by_xpath("//span[text()='Manage']").click()
                    sleep(2)
                    browser.find_element_by_xpath("//a[text()='Licensing']").click()
                    wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div/div/div[1]/div/div[2]/div/div[2]/div[2]/div/div[2]/div/div/div/div[3]/div/div/div[3]/div/div[1]/img")))
                    sleep(2)
                    shot(folder, "license")
                    logging.info("Captured a screenshot of License page")

                except (RuntimeError, TypeError, NameError) as e:
                    logging.error(e)
            

        
        else: logging.warning("{} is not pingable".format(key))

    browser.close()
      
    logging.info("Executed iDrac validation for 3GI project")
    iDRAC(ex, folder, "3GI")

def stargate(ex, folder):

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Stargate")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    # Get mps server value
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Stargate")
        mps = {}
        for i in range(1, sheet.nrows, +1):
            mps[sheet.cell_value(i,1)] = sheet.cell_value(i,4)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    

    #Get Project name from Excel
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Engineer_Info")
        projectName = str(sheet.cell(1,2)).replace("'", "").replace("text:", "").capitalize()
        print(projectName)

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    logging.info("Data from Stargate tab has been retrieved Successfully")

    #Create Stargate directory or not
    try:
        if os.path.isdir('{}\\Stargate Validations'.format(folder)):
            shutil.rmtree('{}\\Stargate Validations'.format(folder))

        folder = "{}\\Stargate Validations".format(folder)
        os.mkdir(folder)
        logging.info("Stargate Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Stargate validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():
        #Checks if server is pingable
        if ping(value, key):
            
            logging.info("{} is pingable".format(key))

            if mps.get(key):
                try:
                    userPath = os.environ['USERPROFILE']
                    p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Stargate {} '{}' {} '{}'".format(userPath, value, folder, key, projectName)])
                    p.communicate()
                    logging.info("run_on_remote.ps1 script has been executed successfully")
                except (RuntimeError, TypeError, NameError) as e:
                    logging.error(e)
        
        else: logging.warning("{} is not pingable".format(key))
        
    logging.info("Executed iDrac validation for Stargate project")
    iDRAC(ex, folder, "Stargate")

def sau(ex, folder): #Sau servers screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("Sau")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    logging.info("Data from Sau tab has been retrieved Successfully")

    #Create Sau directory or not
    try:
        if os.path.isdir('{}\\Sau Validations'.format(folder)):
            shutil.rmtree('{}\\Sau Validations'.format(folder))

        folder = "{}\\Sau Validations".format(folder)
        os.mkdir(folder)
        logging.info("Sau Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Sau validation
    #--------------------------------------------------------------------------------------------------------------------------------------

    for key, value in dic.items():
        
        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create a server directory
            try:
                if os.path.isdir('{}\\{}'.format(folder, key)):
                    shutil.rmtree('{}\\{}'.format(folder, key))

                folder1 = "{}\\{}".format(folder, key)
                os.mkdir(folder1)
                logging.info("A folder for server {} has been created at {}".format(key, folder1))
            except (PermissionError) as e:
                logging.error(e)
                os._exit(1)
            
            #Execute remote powershell script
            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Remote Run\\run_on_remote.ps1' Sau {} '{}'".format(userPath, value, folder1)])
                p.communicate()
                logging.info("run_on_remote.ps1 script has been executed successfully")
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)

        else: logging.warning("{} is not pingable".format(key))
            
    logging.info("Executed iDrac validation for Sau servers")
    #iDRAC(ex, folder, "Sau")

def esxi(ex, folder): #ESXi Screenshots creation

    #Read from excel file
    try:
        wb = xlrd.open_workbook(ex)
        sheet = wb.sheet_by_name("ESXi")
        dic = {}
        for i in range(1, sheet.nrows, +1):
            dic[sheet.cell_value(i,1)] = sheet.cell_value(i,0)
    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    logging.info("Data from ESXi tab has been retrieved Successfully")

    #Create ESXi directory or not
    try:
        if os.path.isdir('{}\\ESXi Validations'.format(folder)):
            shutil.rmtree('{}\\ESXi Validations'.format(folder))

        folder = "{}\\ESXi Validations".format(folder)
        os.mkdir(folder)
        logging.info("ESXi Validations directory has been created Successfully")
    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #ESXi Validations
    #--------------------------------------------------------------------------------------------------------------------------------

    #Set security certificate as trusted
    options = webdriver.ChromeOptions()
    options.add_argument('ignore-certificate-errors')
    options.add_experimental_option("excludeSwitches", ['enable-automation']) #set banner of controlled chrome off
    logging.info("Security certificate has been set Successfuly")

    #create browser element + assign chrome driver
    browser = webdriver.Chrome('Files\\chromedriver.exe', chrome_options=options)
    browser.implicitly_wait(90)

    #Set wait element time
    wait = WebDriverWait(browser, 90)

    for key, value in dic.items():

        #Checks if server is pingable
        if ping(value, key):

            logging.info("{} is pingable".format(key))

            #Create specific ESXi directory
            try:
                if os.path.isdir('{}\\{}'.format(folder, key)):
                    shutil.rmtree('{}\\{}'.format(folder, key))

                folder1 = "{}\\{}".format(folder, key)
                os.mkdir(folder1)
                logging.info("A folder for server {} has been created at {}".format(key, folder1))
            except (PermissionError) as e:
                logging.error(e)
                os._exit(1)

            try:
                browser.get("https://{}/ui/#/host/vms".format(value)) #open browser window

                #login into ESXi web page
                ESXi_username = browser.find_element_by_id('username')
                ESXi_password = browser.find_element_by_id('password')
                ESXi_username.send_keys("root")
                ESXi_password.send_keys("password")
                browser.find_element_by_id("submit").click()

                logging.info("Logged into ESXi web interface Successfully")

                browser.maximize_window() #maximize browser window
                wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='createVMButton']")))
                sleep(2)

                #Take a screenshot of all the vms
                shot(folder1, "vms")
                logging.info("Captured a screenshot of Virtual Machines page")

                #Take a screenshot of the storage span
                browser.find_element_by_xpath("//span[text()='Storage']").click()
                wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='datastoreGrid']/div/div[2]/div/table/thead/tr/th[1]")))
                sleep(2)
                shot(folder1,"storage")
                logging.info("Captured a screenshot of Storage page")

                #Take a screenshot of the network span
                browser.find_element_by_xpath("/html/body/div/div/div[1]/div/div[2]/div/div[2]/div[1]/div/div[2]/div/div/ul/li[4]/div/a/span[2]").click()
                wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='portgroupGrid']/div/div[2]/div/table/thead/tr/th[1]")))
                sleep(2)
                shot(folder1,"network")
                logging.info("Captured a screenshot of Network page")


                #Take a screenshot of the Host span
                browser.find_element_by_xpath("//span[text()='Host']").click()
                wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div/div/div[1]/div/div[2]/div/div[2]/div[2]/div/div[2]/div/div/div[4]/div/div[1]/div/div[1]")))
                sleep(2)
                shot(folder1, "host")
                logging.info("Captured a screenshot of HOst page")

                #Take a screenshot of the Licensing span
                browser.find_element_by_xpath("//span[text()='Manage']").click()
                sleep(2)
                browser.find_element_by_xpath("//a[text()='Licensing']").click()
                wait.until(EC.visibility_of_element_located((By.XPATH, "/html/body/div/div/div[1]/div/div[2]/div/div[2]/div[2]/div/div[2]/div/div/div/div[3]/div/div/div[3]/div/div[1]/img")))
                sleep(2)
                shot(folder1, "license")
                logging.info("Captured a screenshot of License page")

            except (RuntimeError, TimeoutError) as e:
                logging.error(e)

        else: logging.warning("{} is not pingable".format(key))

    browser.close()

    logging.info("Executed iDrac validation for ESXi servers")
    #iDRAC(ex, folder, "ESXi")

def iDRAC(ex, folder, serverType): #iDrac screenshots creation

    #iDRAC Validations
    #--------------------------------------------------------------------------------------------------------------------------------

    #Read from excel file
    try:
        #Checks if this function call is for iDrac only
        if serverType == "iDrac_Only":
            wb = xlrd.open_workbook(ex)
            sheet = wb.sheet_by_name(serverType)
            dic = {}
            for i in range(1, sheet.nrows, +1):
                dic[sheet.cell_value(i,0)] = sheet.cell_value(i,1)

        #Regular data gathering
        else:
            wb = xlrd.open_workbook(ex)
            sheet = wb.sheet_by_name(serverType)
            dic = {}
            for i in range(1, sheet.nrows, +1):
                dic[sheet.cell_value(i,1)] = sheet.cell_value(i,2)

        #Get mps tab value if serverType is 3GI or Stargate
        if serverType == "3GI" or serverType == "Stargate":
            # Get mps server value
                wb = xlrd.open_workbook(ex)
                sheet = wb.sheet_by_name(serverType)
                mps = {}
                for i in range(1, sheet.nrows, +1):
                    mps[sheet.cell_value(i,1)] = sheet.cell_value(i,4)

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)
    
    logging.info("Data from {} tab has been retrieved Successfully".format(serverType))

    if serverType == "iDrac_Only":

        #Deletes fodler if already exists and creates new one
        try:
            if os.path.isdir('{}\\iDrac Validations'.format(folder)):
                shutil.rmtree('{}\\iDrac Validations'.format(folder))
            folder = "{}\\iDrac Validations".format(folder)
            os.mkdir(folder)
            logging.info("Idrac Validations folder has been created Successfully")
        except (PermissionError) as e:
            print(e)
            os._exit(1)

    #Run validations

    #Count variable for 3GI and Stargate servers idrac photos names
    count = 0

    #Set security certificate as trusted
    options = webdriver.ChromeOptions()
    options.add_argument('ignore-certificate-errors')
    options.add_experimental_option("excludeSwitches", ['enable-automation']) #set banner of controlled chrome off
    logging.info("Security certificate has been set Successfuly")

    #create browser element + assign chrome driver
    browser = webdriver.Chrome('Files\\chromedriver.exe', chrome_options=options)
    browser.implicitly_wait(90)

    #Set wait element time
    wait = WebDriverWait(browser, 90)

    for key, value in dic.items():

        #Checks if server is pingable
        if ping(value, key):

            logging.info("{}'s iDrac is pingable".format(key))

            try:
                userPath = os.environ['USERPROFILE']
                p = subprocess.Popen(["PowerShell", "& '{}\\Desktop\\Scripts\\Automated Validations\\Files\\delSession.ps1' {}".format(userPath, value)])
                p.communicate()
                logging.info("Cleaned all sessions from iDrac {}".format(value))
                
            except (RuntimeError, TypeError, NameError) as e:
                logging.error(e)

            #Set folder to put screenshot in
            folder1 = "{}\\{}".format(folder, key)

            if serverType == "iDrac_Only":

                #Deletes fodler if already exists and creates new one
                try:
                    if os.path.isdir('{}\\{}'.format(folder, key)):
                        shutil.rmtree('{}\\{}'.format(folder, key))
                    folder1 = "{}\\{}".format(folder, key)
                    os.mkdir(folder1)  
                    logging.info("A folder for idrac {} has been created at {}".format(key, folder1))
                except (PermissionError) as e:
                    print(e)
                    os._exit(1)

            #Do regular check for all servers
            if serverType != "3GI" and serverType != "Stargate":

                try:

                    if requests.get("https://{}/restgui/start.html?login".format(value), verify=False).status_code == 200:

                        browser.get("https://{}/restgui/start.html?login".format(value)) #open browser window

                        #login into iDrac web page
                        iDRAC_username = browser.find_element_by_name("username")
                        iDRAC_password = browser.find_element_by_name("password")
                        iDRAC_username.send_keys("root")
                        iDRAC_password.send_keys("password")

                        #Wait for button to be clickable
                        wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[2]/idrac-start-screen/div/div/div/div/div/form/div[2]/div[3]/button")))
                        browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("/html/body/div[2]/idrac-start-screen/div/div/div/div/div/form/div[2]/div[3]/button")) 

                        logging.info("Logged into iDrac Successfully")

                        wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='module-div']/div[2]/div/div[2]/div[2]/div/div[2]/table/tbody/tr[6]/td[1]")))
                        browser.maximize_window() #maximize browser window

                        browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='settings']"))
                        wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='module-div']/div[2]/div/div[2]/div[3]/div[2]/table/tbody/tr[1]/td[2]")))
                        sleep(4)
                        shot(folder1, "iDrac")

                        if serverType in "Tbos":
                            browser.execute_script("arguments[0].click();", browser.find_element_by_id("system"))
                            wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='system.inventory']/a")))
                            browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='system.inventory']/a"))
                            wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='system.inventory.fwinventory.fw_inventory']")))
                            browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='system.inventory.fwinventory.fw_inventory']"))
                            wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='module-div']/div[2]/div[1]/div/div[2]/table/tbody/tr[10]/td[1]")))
                            sleep(2)
                            shot(folder1, "iDrac_inventory")
                    else:

                        #create browser element + assign chrome driver
                        browser.get("https://{}/login.html".format(value)) #open browser window

                        #login into iDrac web page
                        username = browser.find_element_by_id("user")
                        password = browser.find_element_by_id("password")
                        username.send_keys('root')
                        password.send_keys('password')

                        wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='btnOK']")))
                        sleep(2)
                        browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='btnOK']"))
                        sleep(10)

                        #Maximise windows size for screenshot
                        browser.maximize_window()

                        #Move to needed frame and press idrac settings anchor tag
                        browser.switch_to.frame("treelist_id")
                        wait.until(EC.presence_of_element_located((By.XPATH, "//*[@id='a_iDRAC Settings ']")))
                        sleep(3)
                        browser.execute_script('document.getElementById("a_iDRAC Settings ").click()')
                        sleep(12)
                        shot(folder1, "iDrac")

                        #Tbos only validations
                        if serverType in "Tbos":
                            browser.find_element_by_id("a_Server").click()
                            browser.switch_to.default_content()
                            browser.execute_script("document.getElementsByName('lsnb').item(0).contentWindow.document.getElementById('lsnb4').click()")
                            browser.execute_script("document.getElementsByName('da').item(0).setAttribute('id', 'scinven')")
                            wait.until(EC.presence_of_element_located((By.ID, "scinven")))
                            browser.switch_to.frame("scinven")
                            sleep(15)
                            browser.execute_script("document.getElementById('jb2').click()")
                            sleep(10)
                            shot(folder1, "iDrac_inventory")


                except (RuntimeError, TimeoutError) as e:
                    logging.error(e)

            #Else doing iDrac check only for HVS and ESXi servers
            else:

                if mps.get(key) != 1:

                    #Count how many servers (for hvs stargate validation) has been checked already (for idrac photo name)
                    count = count + 1

                    try:
                        if requests.get("https://{}/restgui/start.html?login".format(value), verify=False).status_code == 200:
                            browser.get("https://{}/restgui/start.html?login".format(value)) #open browser window

                            # WebDriverWait(browser, 10).until(EC.presence_of_element_located((By.NAME, "username")))

                            #login into iDrac web page
                            iDRAC_username = browser.find_element_by_name("username") #id=user
                            iDRAC_password = browser.find_element_by_name("password") #id=password
                            iDRAC_username.send_keys("root")
                            iDRAC_password.send_keys("password")

                            #Wait for button to be clickable
                            wait.until(EC.element_to_be_clickable((By.XPATH, "/html/body/div[2]/idrac-start-screen/div/div/div/div/div/form/div[2]/div[3]/button")))
                            browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("/html/body/div[2]/idrac-start-screen/div/div/div/div/div/form/div[2]/div[3]/button")) 

                            logging.info("Logged into iDrac Successfully")

                            wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='module-div']/div[2]/div/div[2]/div[2]/div/div[2]/table/tbody/tr[6]/td[1]")))
                            browser.maximize_window() #maximize browser window

                            browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='settings']"))
                            wait.until(EC.visibility_of_element_located((By.XPATH, "//*[@id='module-div']/div[2]/div/div[2]/div[3]/div[2]/table/tbody/tr[1]/td[2]")))
                            sleep(4)
                            shot(folder, "iDrac-{}".format(count))
                        
                        else:

                            #create browser element + assign chrome driver
                            browser.get("https://{}/login.html".format(value)) #open browser window

                            #login into iDrac web page
                            username = browser.find_element_by_id("user")
                            password = browser.find_element_by_id("password")
                            username.send_keys('root')
                            password.send_keys('password')

                            wait.until(EC.element_to_be_clickable((By.XPATH, "//*[@id='btnOK']")))
                            sleep(2)
                            browser.execute_script("arguments[0].click();", browser.find_element_by_xpath("//*[@id='btnOK']"))
                            sleep(10)

                            #Maximise windows size for screenshot
                            browser.maximize_window()

                            #Move to needed frame and press idrac settings anchor tag
                            browser.switch_to.frame("treelist_id")
                            wait.until(EC.presence_of_element_located((By.XPATH, "//*[@id='a_iDRAC Settings ']")))
                            sleep(3)
                            browser.execute_script('document.getElementById("a_iDRAC Settings ").click()')
                            sleep(10)
                            shot(folder, "iDrac-{}".format(count))

                    except (RuntimeError, TimeoutError) as e:
                        logging.error(e)     

        else: logging.warning("The iDrac of {} is not pingable".format(key))
    browser.close()

if __name__ == '__main__':

    #Create a window for files gathering process
    root = Tk()
    root.withdraw()

    #Choose validations folder's location
    folder_selected = filedialog.askdirectory()

    #Choose excel validation file
    root.withdraw()
    excel_file = filedialog.askopenfilename(filetypes =[('Excel Files', '*.xlsx')])

    #Create log file and folder if not already exist
    try:
        if os.path.isdir('{}\\..\\Logs'.format(folder_selected)):
            shutil.rmtree('{}\\..\\Logs'.format(folder_selected))
        os.mkdir('{}\\..\\Logs'.format(folder_selected))
        logging.basicConfig(filename='{}\\..\\Logs\\log.txt'.format(folder_selected), level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

    except (PermissionError) as e:
        logging.error(e)
        os._exit(1)

    #Validate what needed only
    try:
        wb = xlrd.open_workbook(excel_file)
        sheet = wb.sheet_by_name("Validate")
        dic = {}
        for i in range(0, sheet.ncols, +1):
            dic[sheet.cell_value(0,i)] = sheet.cell_value(1,i) #arr.append(sheet.cell_value(i,0))

        logging.info("Collected data from Validate tab successfully")

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    #Dict of servers to create docx files for
    serversDic = {}
    logging.info('Servers Dictionary has been created successfuly')

    #Start validations
    for key, value in dic.items():

        if key == "iDrac_Only":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started iDrac Only validations')
                iDRAC(excel_file, folder_selected, "iDrac_Only")

        if key == "ESXi":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started ESXi validations')
                esxi(excel_file, folder_selected)
                serversDic["ESXi"] = True

            else:
                serversDic["ESXi"] = False

        if key == "3GI":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started 3GI validations')
                threegi(excel_file, folder_selected)
                serversDic["3GI"] = True
            else:
                serversDic["3GI"] = False

        if key == "Stargate":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Stargate validations')
                stargate(excel_file, folder_selected)
                serversDic["Stargate"] = True
            else:
                serversDic["Stargate"] = False

        if key == "Presto":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Presto validations')
                presto(excel_file, folder_selected)
                serversDic["Presto"] = True
            else:
                serversDic["Presto"] = False

        if key == "Sau":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Sau validations')
                sau(excel_file, folder_selected)
                serversDic["Sau"] = True
            else:
                serversDic["Sau"] = False

        if key == "Kafka":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Kafka validations')
                kafka(excel_file, folder_selected)
                serversDic["Kafka"] = True
            else:
                serversDic["Kafka"] = False

        if key == "Workers":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Workers validations')
                workers(excel_file, folder_selected)
                serversDic["Workers"] = True
            else:
                serversDic["Workers"] = False

        if key == "MDE":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started MDE validations')
                mde(excel_file, folder_selected)
                serversDic["MDE"] = True
            else:
                serversDic["MDE"] = False

        if key == "Tbos":
            if value == 1:
                logging.info('-------------------------------------------------------------------------------------------------------------')
                logging.info('Started Tbos validations')
                tbos(excel_file, folder_selected)
                serversDic["Tbos"] = True
            else:
                serversDic["Tbos"] = False

    #Create docx files
    logging.info('-------------------------------------------------------------------------------------------------------------')
    logging.info('Started Creating DOCX Files')
    docxReady(excel_file, folder_selected, serversDic)

    #Delete all pictures inside validation folders created
    logging.info('-------------------------------------------------------------------------------------------------------------')
    logging.info('Deleting all jpg and png files inside {}'.format(folder_selected))

    try: 
        for root, dirs, files in os.walk(r'{}'.format(folder_selected)):
            for fileName in files:
                picture = os.path.join(root, fileName)
                if "iDrac Validations" not in picture:
                    if "jpg" in picture or "png" in picture:
                        os.remove(picture)
                else:
                    logging.info("Skipped iDrac photo in iDrac Only folder")
        logging.info("All files has been deleted from {}".format(folder_selected))

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)


    #Zip all to one file!
    # -------------------------------------------------------------

    logging.info('-------------------------------------------------------------------------------------------------------------')
    logging.info('Zipping "{}" into one file'.format(folder_selected))

    #Search excel file for Engineer data
    try:
        wb = xlrd.open_workbook(excel_file)
        sheet = wb.sheet_by_name("Engineer_Info")
        eng = {}
        for i in range(0, sheet.ncols, +1):
            eng[sheet.cell_value(0,i)] = sheet.cell_value(1,i)
        logging.info("Data from Engineer_Info tab has been retrieved Successfully")

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)    
    
    try:
        shutil.make_archive(os.path.normpath('{}\\..\\{}-Validations'.format(folder_selected, eng.get("Project_Name"))), 'zip', folder_selected)
        logging.info("ZIP Archive was created successfully")

    except (RuntimeError, TypeError, NameError) as e:
        logging.error(e)

    #Open ziped file folder
    shutil.copyfile(os.path.normpath('{}\\..\\{}-Validations.zip'.format(folder_selected, eng.get('Project_Name'))), os.path.normpath('{}\\..\\{}-Validations.zip.txt'.format(folder_selected, eng.get('Project_Name'))))
    os.startfile(os.path.normpath('{}\\..'.format(folder_selected)))
    logging.info('Validations process has been completed')
